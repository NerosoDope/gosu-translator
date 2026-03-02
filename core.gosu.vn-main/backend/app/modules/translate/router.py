import base64
import csv
import io
import json
import logging
import re
import xml.etree.ElementTree as ET
from typing import Any, Dict, List, Optional, Tuple

from fastapi import APIRouter, Depends, File, Form, HTTPException, Query, UploadFile
from fastapi.responses import Response
from sqlalchemy.ext.asyncio import AsyncSession

from app.db.session import get_db
from app.modules.translate.schemas import (
    ExportFileRequest,
    ParseAllRowsResponse,
    ParseFileResponse,
    ProofreadBatchRequest,
    ProofreadBatchResponse,
    ProofreadBatchResultItem,
    ProofreadRowRequest,
    ProofreadRowResponse,
    TranslateFileResponse,
    TranslateRequest,
    TranslateResponse,
)
from app.modules.translate.service import (
    proofread_with_ai,
    proofread_with_ai_batch,
    save_translation_to_cache,
    translate_check_only,
    translate_with_ai_batch,
    translate_with_priority,
    verify_gemini_api_key,
)

logger = logging.getLogger(__name__)
router = APIRouter(prefix="", tags=["translate"])

PREVIEW_ROW_LIMIT = 5
TRANSLATE_FILE_MAX_ROWS = 500
# Số segment tối đa gửi trong 1 lần gọi AI batch (fallback nếu không ước tính token)
TRANSLATE_BATCH_SIZE = 30
# Token input tối đa mỗi batch (Gemini flash ~1M context, nhưng ta giới hạn để kiểm soát cost)
TRANSLATE_BATCH_MAX_INPUT_TOKENS = 1500


def _estimate_tokens(text: str) -> int:
    """Ước tính số token: tiếng Việt/Anh ≈ 3 ký tự/token (không cần tiktoken)."""
    return max(1, len(text) // 3)


def _build_token_batches(
    pending: "List[Tuple[int, str, str]]",
    max_tokens: int = TRANSLATE_BATCH_MAX_INPUT_TOKENS,
    max_items: int = TRANSLATE_BATCH_SIZE,
) -> "List[List[Tuple[int, str, str]]]":
    """
    Gom pending cells thành batches dựa trên ước tính token.
    Mỗi item chiếm: _estimate_tokens(text) + 6 (overhead số dòng + dấu chấm).
    """
    batches: List[List[Tuple[int, str, str]]] = []
    current: List[Tuple[int, str, str]] = []
    current_tokens = 0
    for item in pending:
        tok = _estimate_tokens(item[2]) + 6
        if current and (current_tokens + tok > max_tokens or len(current) >= max_items):
            batches.append(current)
            current = [item]
            current_tokens = tok
        else:
            current.append(item)
            current_tokens += tok
    if current:
        batches.append(current)
    return batches


def _cell_to_str(value) -> str:
    """Chuyển giá trị ô (Excel/CSV) sang chuỗi."""
    if value is None:
        return ""
    if hasattr(value, "isoformat"):  # datetime.date/datetime
        return value.isoformat()
    return str(value).strip()


def _json_value_to_str(value) -> str:
    """Chuyển giá trị từ JSON sang chuỗi: số/chuỗi/None như ô; dict/list serialize thành JSON (khác Excel)."""
    if value is None:
        return ""
    if isinstance(value, (dict, list)):
        try:
            return json.dumps(value, ensure_ascii=False).strip()
        except (TypeError, ValueError):
            return str(value).strip()
    if hasattr(value, "isoformat"):
        return value.isoformat()
    return str(value).strip() if value is not None else ""


async def _parse_excel(content: bytes) -> ParseFileResponse:
    """Đọc file .xlsx (openpyxl), trả về columns và preview_rows.
    Bỏ qua cột trống trong header và hàng mà toàn bộ cell đều trống.
    """
    try:
        import openpyxl
    except ImportError:
        raise HTTPException(
            status_code=503,
            detail="openpyxl chưa được cài đặt. Vui lòng cài: pip install openpyxl",
        )
    workbook = openpyxl.load_workbook(io.BytesIO(content), read_only=True, data_only=True)
    sheet = workbook.active
    if not sheet:
        return ParseFileResponse(columns=[], preview_rows=[])
    # Dòng 1 = header
    header_row = next(sheet.iter_rows(min_row=1, max_row=1, values_only=True), None)
    if not header_row:
        return ParseFileResponse(columns=[], preview_rows=[])

    # Ghi nhận index các cột có tên (bỏ qua cột header trống)
    raw_columns = [_cell_to_str(h) for h in header_row]
    valid_col_indices = [i for i, c in enumerate(raw_columns) if c.strip()]

    # Đảm bảo tên cột unique (trùng thì thêm _2, _3...)
    seen: dict = {}
    columns: List[str] = []
    for i in valid_col_indices:
        c = raw_columns[i]
        key = c
        if key in seen:
            seen[key] += 1
            c = f"{c}_{seen[key]}"
        else:
            seen[key] = 1
        columns.append(c)

    preview_rows = []
    for row in sheet.iter_rows(min_row=2, max_row=1 + PREVIEW_ROW_LIMIT, values_only=True):
        row_dict = {
            col_name: (_cell_to_str(row[orig_i]) if orig_i < len(row) else "")
            for col_name, orig_i in zip(columns, valid_col_indices)
        }
        # Bỏ qua hàng trống (toàn bộ cell đều rỗng)
        if not any(v.strip() for v in row_dict.values()):
            continue
        preview_rows.append(row_dict)
    workbook.close()
    return ParseFileResponse(columns=columns, preview_rows=preview_rows)


async def _parse_csv(content: bytes) -> ParseFileResponse:
    """Đọc file CSV (dòng 1 = header), trả về columns và preview_rows."""
    try:
        text = content.decode("utf-8-sig").strip()
    except Exception:
        text = content.decode("utf-8", errors="replace").strip()
    lines = text.splitlines()
    if not lines:
        return ParseFileResponse(columns=[], preview_rows=[])
    reader = csv.reader(io.StringIO(text))
    header = next(reader, None)
    if not header:
        return ParseFileResponse(columns=[], preview_rows=[])
    columns = [h.strip() or "Cột" for h in header]
    seen = {}
    unique_columns = []
    for c in columns:
        key = c
        if key in seen:
            seen[key] = seen.get(key, 1) + 1
            c = f"{c}_{seen[key]}"
        else:
            seen[key] = 1
        unique_columns.append(c)
    columns = unique_columns
    preview_rows = []
    for _ in range(PREVIEW_ROW_LIMIT):
        row = next(reader, None)
        if row is None:
            break
        row_dict = {}
        for i, col_name in enumerate(columns):
            row_dict[col_name] = (row[i].strip() if i < len(row) else "")
        preview_rows.append(row_dict)
    return ParseFileResponse(columns=columns, preview_rows=preview_rows)


def _rows_from_json(data) -> Tuple[List[str], List[dict]]:
    """Từ cấu trúc JSON (list of dict, dict với key chứa mảng object, dict đơn, list of list, ...) trả về columns và rows."""
    if isinstance(data, dict):
        # Nếu root là object có key chứa mảng object (vd. {"users": [{...}, {...}]}) thì dùng mảng đó làm dòng
        for _k, v in data.items():
            if isinstance(v, list) and len(v) > 0 and all(isinstance(item, dict) for item in v):
                return _rows_from_json(v)
        columns = [str(c) or "Cột" for c in data.keys()]
        if not columns:
            return ["Nội dung"], []
        row = {c: _json_value_to_str(data.get(k)) for k, c in zip(list(data.keys()), columns)}
        return columns, [row]
    if isinstance(data, list):
        if len(data) == 0:
            return ["Nội dung"], []
        first = data[0]
        if isinstance(first, dict):
            all_keys = []
            for row in data:
                if isinstance(row, dict):
                    for k in row:
                        if k not in all_keys:
                            all_keys.append(k)
            columns = [str(k) or "Cột" for k in all_keys]
            if not columns:
                columns = ["Nội dung"]
            seen = {}
            unique = []
            for c in columns:
                key = c
                if key in seen:
                    seen[key] += 1
                    c = f"{c}_{seen[key]}"
                else:
                    seen[key] = 1
                unique.append(c)
            rows = []
            for item in data:
                if not isinstance(item, dict):
                    continue
                row = {}
                for i, col in enumerate(unique):
                    orig_key = all_keys[i] if i < len(all_keys) else col
                    val = item.get(orig_key)
                    row[col] = _json_value_to_str(val)
                rows.append(row)
            return unique, rows
        if isinstance(first, (list, tuple)):
            # Mảng 2 chiều: dòng đầu = header
            columns = [str(v) or f"Cột{i+1}" for i, v in enumerate(first)]
            if not columns:
                columns = ["Nội dung"]
            seen = {}
            unique = []
            for c in columns:
                key = c
                if key in seen:
                    seen[key] += 1
                    c = f"{c}_{seen[key]}"
                else:
                    seen[key] = 1
                unique.append(c)
            rows = []
            for item in data[1:]:
                if not isinstance(item, (list, tuple)):
                    continue
                row = {}
                for i, col in enumerate(unique):
                    val = item[i] if i < len(item) else None
                    row[col] = _json_value_to_str(val)
                rows.append(row)
            return unique, rows
        return ["Nội dung"], [{"Nội dung": _json_value_to_str(first)}]
    return ["Nội dung"], [{"Nội dung": _json_value_to_str(data)}]


def _decode_text(content: bytes) -> str:
    """Thử nhiều encoding để decode bytes thành str."""
    for encoding in ("utf-8-sig", "utf-8", "cp1252", "latin-1"):
        try:
            return content.decode(encoding).strip().lstrip("\ufeff")
        except Exception:
            continue
    return content.decode("utf-8", errors="replace").strip().lstrip("\ufeff")


# ─────────────────────────────────────────────────────────────────────────────
# Smart Filter — phân loại 5 nhóm:
#   🟢 Natural language   → dịch
#   🟡 Mixed text         → bóc placeholder rồi dịch phần text
#   🔵 Placeholder only   → không dịch
#   🔴 Code-like          → không dịch
#   ⚫ Pure technical     → không dịch
# ─────────────────────────────────────────────────────────────────────────────

# Nhóm ⚫ Pure technical
_RE_URL         = re.compile(r"^https?://\S+$", re.IGNORECASE)
_RE_EMAIL       = re.compile(r"^[^\s@]+@[^\s@]+\.[^\s@]+$")
_RE_ISO_DATE    = re.compile(r"^\d{4}-\d{2}-\d{2}([T ]\d{2}:\d{2}(:\d{2})?)?$")
_RE_PURE_NUMBER = re.compile(r"^-?\d+([.,]\d+)?([eE][+-]?\d+)?$")
_RE_VERSION     = re.compile(r"^v?\d+\.\d+[\d.\-a-zA-Z]*$")
_RE_ISO_LANG    = re.compile(r"^[a-z]{2,3}$")   # en, vi, fr, zh-CN (chuỗi ngắn toàn chữ thường)
_RE_BOOLEAN     = re.compile(r"^(true|false|null|undefined|none|nan)$", re.IGNORECASE)
_RE_FILE_PATH   = re.compile(
    r"^(/|\.\.?/)[\w.\-/]+"
    r"\.(png|jpg|jpeg|gif|webp|svg|json|js|ts|tsx|jsx|css|html|htm|py|rb|go|java|php|vue|xml|yaml|yml|sh|env|conf|ini)$",
    re.IGNORECASE,
)

# Nhóm 🔴 Code-like
_RE_ALL_CAPS    = re.compile(r"^[A-Z][A-Z0-9_]*$")         # USER_LOGIN_SUCCESS, API, CONFIG_KEY
_RE_SNAKE_CASE  = re.compile(r"^[a-z][a-z0-9]*(_[a-z0-9]+)+$")  # user_id, config_key (bắt buộc có dấu _)
_RE_CAMEL_CASE  = re.compile(r"^[a-z]+([A-Z][a-z0-9]+)+$")       # userId, configKey

# Natural language helpers
_RE_HAS_LETTER  = re.compile(r"[a-zA-ZÀ-ỹ\u00C0-\u024F\u1E00-\u1EFF]")
_RE_NATURAL_START = re.compile(r"^[A-ZÀ-Ỵ\u00C0-\u024F\u1E00-\u1EFF][a-zà-ỹ\u00C0-\u024F\u1E00-\u1EFF]")
_RE_PUNCTUATION = re.compile(r"[.,!?;:]")

# Nhóm 🔵 Placeholder — tất cả dạng biến: {x}, {{x}}, %s, %d, %1$s, ${x}, <x>, @x@
_RE_PLACEHOLDER = re.compile(
    r"("
    r"\{\{[^}]+\}\}"           # {{count}}
    r"|\{[^}]+\}"              # {username}
    r"|\$\{[^}]+\}"            # ${var}
    r"|%\d*\$?[sdifcq%]"       # %s %d %1$s %02d
    r"|%[sdicfqoxXeEgGp%]"     # C-style printf
    r"|@[A-Za-z_][A-Za-z0-9_]*@"  # @name@
    r"|<[A-Za-z_][A-Za-z0-9_]*>"  # <var>
    r")"
)


def _should_translate_string(s: str, smart_filter: bool) -> bool:
    """
    Phân loại chuỗi thành 5 nhóm. Khi smart_filter=True áp dụng toàn bộ rule.
    Khi smart_filter=False chỉ loại bỏ chuỗi rỗng và không có chữ cái.

    Nhóm 🟡 Mixed text (có cả text + placeholder): _extract_placeholders xử lý
    trước khi gọi hàm này — hàm chỉ nhận chuỗi đã thay __PH_n__.
    """
    if not s or not isinstance(s, str):
        return False
    s = s.strip()
    if not s:
        return False

    if not smart_filter:
        # Chế độ tắt smart filter: vẫn bỏ chuỗi hoàn toàn không có chữ cái
        return bool(_RE_HAS_LETTER.search(s)) and len(s) > 1

    # ── BƯỚC 1: Loại bỏ ngay ──────────────────────────────────────────────────

    # ⚫ Chỉ số (kể cả số thập phân, số âm, scientific)
    if _RE_PURE_NUMBER.match(s):
        return False

    # ⚫ Boolean / null
    if _RE_BOOLEAN.match(s):
        return False

    # ⚫ URL
    if _RE_URL.match(s):
        return False

    # ⚫ Email
    if _RE_EMAIL.match(s):
        return False

    # ⚫ File path (bắt đầu bằng / hoặc ./ hoặc ../)
    if _RE_FILE_PATH.match(s):
        return False

    # ⚫ ISO date/datetime
    if _RE_ISO_DATE.match(s):
        return False

    # ⚫ Version string: v1.2.3, 1.0.0-beta
    if _RE_VERSION.match(s):
        return False

    # ⚫ ISO language code: en, vi, fr, zh (2-3 ký tự toàn thường)
    if _RE_ISO_LANG.match(s):
        return False

    # 🔴 ALL_CAPS_CODE: USER_LOGIN_SUCCESS, API_TIMEOUT, CONFIG_KEY, ADMIN
    if _RE_ALL_CAPS.match(s):
        return False

    # 🔴 snake_case_key: user_id, config_key (bắt buộc có ít nhất 1 dấu _)
    if _RE_SNAKE_CASE.match(s):
        return False

    # 🔴 camelCaseKey: userId, configKey
    if _RE_CAMEL_CASE.match(s):
        return False

    # 🔵 Placeholder-only: sau khi bóc hết placeholder không còn gì
    s_no_ph = _RE_PLACEHOLDER.sub("", s).strip()
    if not s_no_ph:
        return False

    # ── BƯỚC 2: Kiểm tra dấu hiệu ngôn ngữ tự nhiên ─────────────────────────
    if not _RE_HAS_LETTER.search(s):
        return False
    if len(s) <= 1:
        return False

    # 🟢 Có khoảng trắng → rất có thể là câu/cụm từ tự nhiên
    if " " in s:
        return True

    # 🟢 Bắt đầu hoa + chữ thường (kể cả Unicode/tiếng Việt)
    if _RE_NATURAL_START.match(s):
        return True

    # 🟢 Có dấu câu
    if _RE_PUNCTUATION.search(s):
        return True

    # 🟢 Dài > 3 ký tự và có chữ cái (final fallback)
    if len(s) > 3 and _RE_HAS_LETTER.search(s):
        return True

    return False


def _extract_placeholders(s: str) -> Tuple[str, List[str]]:
    """
    Bóc placeholder khỏi chuỗi, thay bằng __PH_n__.
    Trả về (chuỗi đã thay, danh sách placeholder gốc theo thứ tự).
    Hỗ trợ: {x}, {{x}}, ${x}, %s/%d/%1$s, @name@, <var>.
    """
    if not s:
        return s, []
    placeholders: List[str] = []

    def repl(m: re.Match) -> str:
        placeholders.append(m.group(0))
        return f"__PH_{len(placeholders) - 1}__"

    out = _RE_PLACEHOLDER.sub(repl, s)
    return out, placeholders


def _restore_placeholders(s: str, placeholders: List[str]) -> str:
    """Khôi phục __PH_n__ về placeholder gốc sau khi AI dịch xong."""
    if not placeholders:
        return s
    for i, ph in enumerate(placeholders):
        s = s.replace(f"__PH_{i}__", ph)
    return s


async def _translate_json_recursive(
    obj: Any,
    db: AsyncSession,
    source_lang: str,
    target_lang: str,
    prompt_id: Optional[int],
    context: Optional[str],
    style: Optional[str],
    smart_filter: bool,
    translate_keys: bool,
    game_id: Optional[int] = None,
    game_category_id: Optional[int] = None,
) -> Any:
    """
    Duyệt đệ quy JSON: chỉ dịch value là string (và optionally key).
    Giữ nguyên: key (nếu không bật translate_keys), số, boolean, null, cấu trúc.
    """
    if obj is None:
        return None
    if isinstance(obj, bool):
        return obj
    if isinstance(obj, (int, float)):
        return obj
    if isinstance(obj, str):
        if _should_translate_string(obj, smart_filter):
            try:
                return await translate_with_priority(
                    db,
                    text=obj,
                    source_lang=source_lang,
                    target_lang=target_lang,
                    prompt_id=prompt_id,
                    context=context,
                    style=style,
                    game_id=game_id,
                    game_category_id=game_category_id,
                ) or obj
            except Exception as e:
                logger.warning("Translate JSON string: %s", e)
                return obj
        return obj
    if isinstance(obj, list):
        return [
            await _translate_json_recursive(
                item, db, source_lang, target_lang,
                prompt_id, context, style, smart_filter, translate_keys,
                game_id=game_id, game_category_id=game_category_id,
            )
            for item in obj
        ]
    if isinstance(obj, dict):
        out = {}
        for k, v in obj.items():
            key = k
            if translate_keys and isinstance(k, str) and _should_translate_string(k, smart_filter):
                try:
                    key = await translate_with_priority(
                        db, text=k, source_lang=source_lang, target_lang=target_lang,
                        prompt_id=prompt_id, context=context, style=style,
                        game_id=game_id, game_category_id=game_category_id,
                    ) or k
                except Exception:
                    pass
            out[key] = await _translate_json_recursive(
                v, db, source_lang, target_lang,
                prompt_id, context, style, smart_filter, translate_keys,
                game_id=game_id, game_category_id=game_category_id,
            )
        return out
    return obj


def _plain_text_lines_to_rows(text: str) -> Tuple[List[str], List[dict]]:
    """Khi file .json không phải JSON hợp lệ: coi mỗi dòng là một dòng dữ liệu, cột 'Nội dung'."""
    lines = [line.strip() for line in text.splitlines() if line.strip()]
    return ["Nội dung"], [{"Nội dung": line} for line in lines]


async def _parse_json(content: bytes) -> ParseFileResponse:
    """Đọc file JSON (array of objects, object, array of arrays, hoặc array rỗng). Nếu không phải JSON hợp lệ thì đọc theo từng dòng (plain text)."""
    if not content:
        return ParseFileResponse(columns=["Nội dung"], preview_rows=[])
    text = _decode_text(content)
    if not text:
        return ParseFileResponse(columns=["Nội dung"], preview_rows=[])
    try:
        data = json.loads(text)
    except json.JSONDecodeError:
        columns, all_rows = _plain_text_lines_to_rows(text)
        preview_rows = all_rows[:PREVIEW_ROW_LIMIT]
        return ParseFileResponse(columns=columns, preview_rows=preview_rows)
    columns, all_rows = _rows_from_json(data)
    if not columns:
        columns = ["Nội dung"]
    preview_rows = all_rows[:PREVIEW_ROW_LIMIT]
    return ParseFileResponse(columns=columns, preview_rows=preview_rows)


def _read_full_json(content: bytes, max_rows: int) -> Tuple[List[str], List[dict]]:
    """Đọc toàn bộ file JSON, trả về columns và rows (tối đa max_rows). Nếu không phải JSON hợp lệ thì đọc theo từng dòng."""
    if not content:
        return [], []
    text = _decode_text(content)
    if not text:
        return [], []
    try:
        data = json.loads(text)
    except json.JSONDecodeError:
        columns, all_rows = _plain_text_lines_to_rows(text)
        return columns, all_rows[:max_rows]
    columns, all_rows = _rows_from_json(data)
    if not columns:
        columns = ["Nội dung"]
    return columns, all_rows[:max_rows]


def _xml_local_tag(el: ET.Element) -> str:
    """Lấy tên thẻ không namespace (phần sau dấu })."""
    tag = el.tag if hasattr(el, "tag") else ""
    if isinstance(tag, str) and "}" in tag:
        return tag.split("}", 1)[1]
    return tag or "col"


def _xml_element_full_text(el: ET.Element) -> str:
    """Lấy toàn bộ nội dung text của phần tử XML (text + nội dung con lồng nhau + tail). Khác Excel: XML có cấu trúc cây."""
    if el is None:
        return ""
    parts = [el.text or ""]
    for child in el:
        parts.append(_xml_element_full_text(child))
        parts.append(child.tail or "")
    return "".join(parts).strip()


def _rows_from_android_xml(root: ET.Element) -> Tuple[List[str], List[dict]]:
    """
    Parser cho Android Resource XML (<resources>).
    Columns: ["name", "value"]
    - <string name="X">text</string>               → {name:"X", value:"text"}
    - <plurals name="X"><item qty="one">t</item>   → {name:"X[one]", value:"t"}
    - <string-array name="X"><item>t</item>        → {name:"X[0]", value:"t"}
    Bỏ qua translatable="false". Tự động strip CDATA markers nếu có.
    """
    columns = ["name", "value"]
    rows: List[dict] = []
    for el in root:
        tag = _xml_local_tag(el)
        if el.attrib.get("translatable", "true").lower() == "false":
            continue
        el_name = el.attrib.get("name", "")
        if tag == "string":
            value = _xml_element_full_text(el).strip()
            rows.append({"name": el_name, "value": value})
        elif tag == "plurals":
            for item in el:
                qty = item.attrib.get("quantity", "")
                value = _xml_element_full_text(item).strip()
                rows.append({"name": f"{el_name}[{qty}]", "value": value})
        elif tag == "string-array":
            for idx, item in enumerate(el):
                value = _xml_element_full_text(item).strip()
                rows.append({"name": f"{el_name}[{idx}]", "value": value})
        elif tag in ("integer", "bool", "color", "dimen"):
            # Giá trị kỹ thuật — bỏ qua, không dịch
            continue
        else:
            # Tag không xác định nhưng có text
            value = _xml_element_full_text(el).strip()
            if value:
                rows.append({"name": el_name or tag, "value": value})
    return columns, rows


def _rows_from_xliff(root: ET.Element) -> Tuple[List[str], List[dict]]:
    """
    Parser cho XLIFF 1.2/2.0 (<xliff>, <file>).
    Columns: ["id", "source", "target"]
    """
    columns = ["id", "source", "target"]
    rows: List[dict] = []

    def _collect_units(node: ET.Element) -> None:
        tag = _xml_local_tag(node)
        if tag in ("trans-unit", "unit"):
            uid = node.attrib.get("id", "")
            src_el = node.find(".//{*}source") or node.find("source")
            tgt_el = node.find(".//{*}target") or node.find("target")
            src = _xml_element_full_text(src_el).strip() if src_el is not None else ""
            tgt = _xml_element_full_text(tgt_el).strip() if tgt_el is not None else ""
            if src or tgt:
                rows.append({"id": uid, "source": src, "target": tgt})
        for child in node:
            _collect_units(child)

    _collect_units(root)
    return columns, rows


def _rows_from_xml(root: ET.Element) -> Tuple[List[str], List[dict]]:
    """
    Dispatcher: nhận diện định dạng XML rồi chọn parser phù hợp.
    - Android Resources (<resources>)  → _rows_from_android_xml
    - XLIFF (<xliff>/<file>)           → _rows_from_xliff
    - Generic repeating-container XML  → logic cũ
    """
    root_tag = _xml_local_tag(root)

    # ── Android Resource XML ──────────────────────────────────────────────────
    if root_tag == "resources":
        return _rows_from_android_xml(root)

    # ── XLIFF ─────────────────────────────────────────────────────────────────
    if root_tag in ("xliff", "file") or root.find(".//{*}trans-unit") is not None or root.find(".//{*}unit") is not None:
        return _rows_from_xliff(root)

    # ── Generic: tìm container có phần tử con lặp lại ────────────────────────
    children = list(root)
    if not children:
        text = _xml_element_full_text(root)
        return (["Nội dung"], [{"Nội dung": text}]) if text else (["Nội dung"], [])

    row_container = None
    for el in children:
        sub = list(el)
        if len(sub) > 1 and (row_container is None or len(sub) > len(list(row_container))):
            row_container = el

    if row_container is not None:
        row_elements = list(row_container)
        if row_elements:
            first_row = row_elements[0]
            attr_names = list(first_row.attrib.keys())
            child_tags: List[str] = []
            for ch in first_row:
                name = _xml_local_tag(ch)
                if name not in child_tags:
                    child_tags.append(name)
            cols_raw = attr_names + child_tags or ["Nội dung"]
            seen: dict = {}
            unique: List[str] = []
            for c in cols_raw:
                if c in seen:
                    seen[c] += 1
                    unique.append(f"{c}_{seen[c]}")
                else:
                    seen[c] = 1
                    unique.append(c)
            rows = []
            for el in row_elements:
                row_dict: dict = {}
                for i, col in enumerate(unique):
                    if i < len(attr_names):
                        row_dict[col] = (el.attrib.get(attr_names[i], "") or "").strip()
                    else:
                        tag_idx = i - len(attr_names)
                        tag_name = child_tags[tag_idx] if tag_idx < len(child_tags) else col
                        child_el = next((e for e in el if _xml_local_tag(e) == tag_name), None)
                        row_dict[col] = _xml_element_full_text(child_el) if child_el is not None else ""
                rows.append(row_dict)
            return unique, rows

    # Fallback: direct children là rows
    first = children[0]
    sub = list(first)
    if not sub:
        # Mỗi thẻ con = 1 row với cột là tag name, giá trị là text
        all_tags = sorted({_xml_local_tag(e) for e in children})
        if len(all_tags) == 1:
            cols = [all_tags[0]]
            return cols, [{all_tags[0]: _xml_element_full_text(e)} for e in children]
        # Nhiều tag khác nhau → dùng cặp key/value
        return ["tag", "value"], [{"tag": _xml_local_tag(e), "value": _xml_element_full_text(e)} for e in children]

    cols_set: List[str] = []
    for c in sub:
        nm = _xml_local_tag(c)
        if nm not in cols_set:
            cols_set.append(nm)
    seen2: dict = {}
    unique2: List[str] = []
    for c in cols_set:
        if c in seen2:
            seen2[c] += 1
            unique2.append(f"{c}_{seen2[c]}")
        else:
            seen2[c] = 1
            unique2.append(c)
    rows2 = []
    for el in children:
        row: dict = {}
        for i, col in enumerate(unique2):
            orig = cols_set[i] if i < len(cols_set) else col
            child_el = next((e for e in el if _xml_local_tag(e) == orig), None) or el.find(orig)
            row[col] = _xml_element_full_text(child_el) if child_el is not None else ""
        rows2.append(row)
    return unique2, rows2


def _strip_xml_declaration(text: str) -> str:
    """Bỏ dòng khai báo <?xml ...?> để tránh lỗi parse do encoding."""
    stripped = text.strip().lstrip("\ufeff")
    if stripped.startswith("<?xml"):
        idx = stripped.find("?>")
        if idx != -1:
            stripped = stripped[idx + 2 :].strip()
    return stripped


def _split_xml_declaration(text: str) -> Tuple[str, str]:
    """Trả về (declaration_line, rest). Declaration có thể rỗng."""
    stripped = text.strip().lstrip("\ufeff")
    if stripped.startswith("<?xml"):
        idx = stripped.find("?>")
        if idx != -1:
            return stripped[: idx + 2].strip(), stripped[idx + 2 :].strip()
    return "", stripped


# --- Dịch XML giữ cấu trúc: chỉ text trong thẻ, không dịch attribute/tên thẻ, giữ placeholder ---
XML_UI_TAGS = frozenset({"string", "text", "description", "title", "message", "label", "subtitle", "content"})

# Các thẻ Android Resource chứa <item> con có text cần dịch (plurals, string-array, string).
# Khi parent là một trong các thẻ này, <item> con sẽ được force-translatable.
ANDROID_ITEM_PARENTS = frozenset({"plurals", "string-array", "string"})


def _is_xml_translatable(el: ET.Element, respect_translatable: bool, force: bool = False) -> bool:
    """Chỉ dịch thẻ UI (string, text, ...) hoặc có translatable=\"true\"; bỏ qua translatable=\"false\".
    Nếu force=True (parent là Android container), luôn dịch trừ khi có translatable=\"false\".
    """
    if respect_translatable:
        trans = (el.attrib.get("translatable") or "").strip().lower()
        if trans == "false":
            return False
        if trans == "true":
            return True
    if force:
        return True
    tag = _xml_local_tag(el)
    return tag.lower() in XML_UI_TAGS


async def _translate_xml_recursive(
    el: ET.Element,
    db: AsyncSession,
    source_lang: str,
    target_lang: str,
    prompt_id: Optional[int],
    context: Optional[str],
    style: Optional[str],
    preserve_placeholders: bool,
    respect_translatable: bool,
    smart_filter: bool,
    _force_children: bool = False,
    game_id: Optional[int] = None,
    game_category_id: Optional[int] = None,
) -> None:
    """
    Duyệt cây XML, chỉ dịch text bên trong thẻ (el.text, child.tail). Không dịch attribute, tên thẻ.
    Giữ placeholder {x}, %d, %s, {{count}}. Chỉ dịch thẻ UI hoặc translatable=\"true\".
    _force_children=True khi parent là Android container (plurals, string-array) — buộc dịch <item>.
    Sửa tại chỗ (in-place).
    """
    if el is None:
        return
    translatable = _is_xml_translatable(el, respect_translatable, force=_force_children)

    async def do_translate(t: str) -> str:
        if not t or not t.strip():
            return t
        orig = t
        if preserve_placeholders:
            t, phs = _extract_placeholders(t)
        else:
            phs = []
        if not _should_translate_string(t, smart_filter):
            return orig
        try:
            out = await translate_with_priority(
                db, text=t.strip(), source_lang=source_lang, target_lang=target_lang,
                prompt_id=prompt_id, context=context, style=style,
                game_id=game_id, game_category_id=game_category_id,
            ) or t.strip()
            if preserve_placeholders and phs:
                out = _restore_placeholders(out, phs)
            return out
        except Exception as e:
            logger.warning("Translate XML text: %s", e)
            return orig

    if translatable and el.text and el.text.strip():
        el.text = await do_translate(el.text)

    # Nếu thẻ hiện tại là Android container, children (<item>) phải được force-translate
    current_tag = _xml_local_tag(el).lower()
    force_for_children = current_tag in ANDROID_ITEM_PARENTS

    for child in el:
        await _translate_xml_recursive(
            child, db, source_lang, target_lang,
            prompt_id, context, style,
            preserve_placeholders, respect_translatable, smart_filter,
            _force_children=force_for_children,
            game_id=game_id, game_category_id=game_category_id,
        )
        if translatable and child.tail and child.tail.strip():
            child.tail = await do_translate(child.tail)


def _xml_to_string(root: ET.Element, declaration: str = "") -> str:
    """Serialize element tree thành chuỗi XML. Nếu declaration có giá trị thì ghi lên trước."""
    buf = io.BytesIO()
    if declaration:
        buf.write(declaration.strip().encode("utf-8"))
        buf.write(b"\n")
    ET.ElementTree(root).write(buf, encoding="utf-8", xml_declaration=False, default_namespace="", method="xml")
    return buf.getvalue().decode("utf-8")


async def _parse_xml(content: bytes) -> ParseFileResponse:
    """Đọc file XML. Không trả 400 khi lỗi parse, trả dữ liệu mặc định."""
    if not content:
        return ParseFileResponse(columns=["Nội dung"], preview_rows=[])
    text = _decode_text(content)
    if not text:
        return ParseFileResponse(columns=["Nội dung"], preview_rows=[])
    text = _strip_xml_declaration(text)
    if not text:
        return ParseFileResponse(columns=["Nội dung"], preview_rows=[])
    try:
        root = ET.fromstring(text)
    except ET.ParseError:
        return ParseFileResponse(columns=["Nội dung"], preview_rows=[])
    columns, all_rows = _rows_from_xml(root)
    if not columns:
        columns = ["Nội dung"]
    preview_rows = all_rows[:PREVIEW_ROW_LIMIT]
    return ParseFileResponse(columns=columns, preview_rows=preview_rows)


def _read_full_xml(content: bytes, max_rows: int) -> Tuple[List[str], List[dict]]:
    """Đọc toàn bộ file XML, trả về columns và rows (tối đa max_rows)."""
    if not content:
        return [], []
    text = _decode_text(content)
    if not text:
        return [], []
    text = _strip_xml_declaration(text)
    if not text:
        return [], []
    try:
        root = ET.fromstring(text)
    except ET.ParseError:
        return [], []
    columns, all_rows = _rows_from_xml(root)
    if not columns:
        columns = ["Nội dung"]
    return columns, all_rows[:max_rows]


def _iter_docx_body_items(doc: Any):
    """Yield Paragraph hoặc Table objects theo đúng thứ tự xuất hiện trong document body.
    Khác doc.paragraphs (chỉ trả paragraph ngoài table) và doc.tables (chỉ trả table).
    """
    try:
        from docx.text.paragraph import Paragraph as _Para
        from docx.table import Table as _Tbl
    except ImportError:
        return
    for child in doc.element.body.iterchildren():
        tag = child.tag.split("}")[-1] if "}" in child.tag else child.tag
        if tag == "p":
            yield _Para(child, doc)
        elif tag == "tbl":
            yield _Tbl(child, doc)



def _collect_docx_blocks(doc: Any, max_items: int = 100_000) -> List[Tuple[str, str, Any]]:
    """Thu thập tất cả text blocks theo thứ tự document: paragraph + table cells.

    Returns list of (text, kind, ref):
        kind = 'para'  → ref là Paragraph object
        kind = 'cell'  → ref là Cell object
    Bỏ qua paragraph/cell rỗng. Bỏ qua cell trùng do merged cells.
    """
    try:
        from docx.text.paragraph import Paragraph as _Para
        from docx.table import Table as _Tbl
    except ImportError:
        return []
    blocks: List[Tuple[str, str, Any]] = []
    for item in _iter_docx_body_items(doc):
        if len(blocks) >= max_items:
            break
        if isinstance(item, _Para):
            text = (item.text or "").strip()
            if text:
                blocks.append((text, "para", item))
        elif isinstance(item, _Tbl):
            seen: set = set()
            for row in item.rows:
                for cell in row.cells:
                    cid = id(cell._tc)
                    if cid in seen:
                        continue
                    seen.add(cid)
                    text = (cell.text or "").strip()
                    if text:
                        blocks.append((text, "cell", cell))
    return blocks


def _para_to_html(para: Any) -> str:
    """Chuyển 1 paragraph python-docx thành HTML, giữ heading/bold/italic/list/underline."""
    import html as _html_mod

    style_name: str = (para.style.name or "") if para.style else ""
    # Build inner HTML từ runs
    inner = ""
    for run in para.runs:
        text = run.text or ""
        if not text:
            continue
        text = _html_mod.escape(text)
        if run.bold and run.italic:
            text = f"<strong><em>{text}</em></strong>"
        elif run.bold:
            text = f"<strong>{text}</strong>"
        elif run.italic:
            text = f"<em>{text}</em>"
        if getattr(run, "underline", False):
            text = f"<u>{text}</u>"
        inner += text

    if not inner.strip():
        return ""

    sn_lower = style_name.lower()
    if "heading 1" in sn_lower:
        return f"<h1>{inner}</h1>"
    if "heading 2" in sn_lower:
        return f"<h2>{inner}</h2>"
    if "heading 3" in sn_lower or "heading 4" in sn_lower:
        return f"<h3>{inner}</h3>"
    if "list" in sn_lower or sn_lower.startswith("list"):
        return f"<li>{inner}</li>"
    return f"<p>{inner}</p>"


def _table_to_html(table: Any) -> str:
    """Chuyển 1 python-docx Table thành HTML <table> với border."""
    import html as _html_mod
    rows_html = ""
    for row_idx, row in enumerate(table.rows):
        seen: set = set()
        cells_html = ""
        for cell in row.cells:
            cid = id(cell._tc)
            if cid in seen:
                continue
            seen.add(cid)
            # Lấy inner HTML của các paragraph trong cell (giữ bold/italic)
            cell_inner = ""
            for p in cell.paragraphs:
                ph = _para_to_html(p)
                if ph:
                    # Bỏ tag bọc ngoài (<p>...</p>, <h1>...</h1>...), chỉ lấy nội dung
                    import re as _re
                    m = _re.match(r"^<[^>]+>(.*)</[^>]+>$", ph, _re.DOTALL)
                    cell_inner += (m.group(1) if m else _html_mod.escape(cell.text)) + " "
            cell_inner = cell_inner.strip() or _html_mod.escape(cell.text or "")
            tag = "th" if row_idx == 0 else "td"
            cells_html += f"<{tag} style='border:1px solid #d1d5db;padding:6px 10px;text-align:left'>{cell_inner}</{tag}>"
        rows_html += f"<tr>{cells_html}</tr>"
    return (
        "<div style='overflow-x:auto;margin:8px 0'>"
        "<table style='border-collapse:collapse;width:100%;font-size:13px'>"
        f"{rows_html}"
        "</table></div>"
    )


def _docx_to_preview_html(doc: Any, max_items: int = 120) -> str:
    """Chuyển toàn bộ document (paragraphs + tables) thành HTML theo thứ tự body.
    Paragraphs → giữ heading/bold/italic/list. Tables → <table> với border.
    """
    try:
        from docx.text.paragraph import Paragraph as _Para
        from docx.table import Table as _Tbl
    except ImportError:
        # Fallback: chỉ lấy paragraphs
        parts = []
        for p in list(doc.paragraphs)[:max_items]:
            h = _para_to_html(p)
            if h:
                parts.append(h)
        return "".join(parts)

    parts: list = []
    count = 0
    for item in _iter_docx_body_items(doc):
        if count >= max_items:
            break
        if isinstance(item, _Para):
            h = _para_to_html(item)
            if h:
                parts.append(h)
                count += 1
        elif isinstance(item, _Tbl):
            parts.append(_table_to_html(item))
            count += 1

    # Gom các <li> liên tiếp vào <ul>
    result = ""
    i = 0
    while i < len(parts):
        if parts[i].startswith("<li>"):
            result += "<ul>"
            while i < len(parts) and parts[i].startswith("<li>"):
                result += parts[i]
                i += 1
            result += "</ul>"
        else:
            result += parts[i]
            i += 1
    return result


async def _parse_docx(content: bytes) -> ParseFileResponse:
    """Đọc file .docx, trả về columns và preview_rows.
    Luôn thu thập tất cả nội dung (paragraphs + table cells) theo thứ tự document.
    """
    try:
        from docx import Document
    except ImportError:
        raise HTTPException(
            status_code=503,
            detail="Hỗ trợ DOCX cần cài thư viện: pip install python-docx (sau đó khởi động lại backend).",
        )
    if not content or len(content) < 4:
        raise HTTPException(status_code=400, detail="File DOCX rỗng hoặc không đúng định dạng.")
    try:
        doc = Document(io.BytesIO(content))
    except Exception as e:
        raise HTTPException(status_code=400, detail=f"Không đọc được file DOCX (file có thể hỏng hoặc không phải .docx): {str(e)[:200]}")

    blocks = _collect_docx_blocks(doc, max_items=PREVIEW_ROW_LIMIT)
    columns = ["Nội dung"]
    preview_rows = [{"Nội dung": text} for text, _kind, _ref in blocks]
    preview_html = _docx_to_preview_html(doc)
    return ParseFileResponse(columns=columns, preview_rows=preview_rows, preview_html=preview_html)


def _read_full_docx(content: bytes, max_rows: int) -> Tuple[List[str], List[dict]]:
    """Đọc toàn bộ file .docx, trả về columns và rows (tối đa max_rows).
    Luôn thu thập tất cả nội dung (paragraphs + table cells) theo thứ tự document.
    """
    try:
        from docx import Document
    except ImportError:
        raise HTTPException(
            status_code=503,
            detail="Hỗ trợ DOCX cần cài: pip install python-docx và khởi động lại backend.",
        )
    if not content or len(content) < 4:
        return [], []
    try:
        doc = Document(io.BytesIO(content))
    except Exception:
        return [], []

    blocks = _collect_docx_blocks(doc, max_items=max_rows)
    rows = [{"Nội dung": text} for text, _kind, _ref in blocks]
    return ["Nội dung"], rows


def _replace_para_text_keep_fmt(para: Any, new_text: str) -> None:
    """Thay text của paragraph mà vẫn giữ nguyên toàn bộ formatting (bold/italic/underline/font/size/color...).

    Chiến lược:
    - 1 run → cập nhật w:t trực tiếp, giữ nguyên rPr.
    - Nhiều runs, cùng format → gom vào run[0], xóa text các run còn lại.
    - Nhiều runs, format khác nhau → phân phối text dịch thông minh:
        1. Thử tìm text gốc của từng run trong text dịch (verbatim, case-insensitive)
           để giữ đúng vị trí bold/italic.
        2. Nếu không tìm được → phân phối tỉ lệ theo độ dài gốc, giữ ranh giới từ.
    """
    try:
        from docx.oxml.ns import qn
        from lxml import etree
    except ImportError:
        para.text = new_text
        return

    XML_SPACE = "{http://www.w3.org/XML/1998/namespace}space"

    def _set_run_text(r: Any, txt: str) -> None:
        t = r._r.find(qn("w:t"))
        if t is None:
            t = etree.SubElement(r._r, qn("w:t"))
        t.text = txt
        t.set(XML_SPACE, "preserve")

    def _fmt_key(r: Any) -> str:
        rpr = r._r.find(qn("w:rPr"))
        return etree.tostring(rpr, encoding="unicode") if rpr is not None else ""

    runs = para.runs
    if not runs:
        para.add_run(new_text)
        return

    if len(runs) == 1:
        _set_run_text(runs[0], new_text)
        return

    # Kiểm tra tất cả runs có cùng format không
    fmt_keys = [_fmt_key(r) for r in runs]
    if len(set(fmt_keys)) == 1:
        # Tất cả cùng format → gom vào run[0], clear các run còn lại
        _set_run_text(runs[0], new_text)
        for run in runs[1:]:
            _set_run_text(run, "")
        return

    # ── Mixed formatting ──────────────────────────────────────────────────────
    # Bước 1: thử tìm text gốc của mỗi run trong text dịch (verbatim)
    orig_texts = [r.text or "" for r in runs]
    distributed: list[str] = [""] * len(runs)
    remaining = new_text

    matched_any = False
    for i, orig in enumerate(orig_texts):
        stripped = orig.strip()
        if not stripped:
            continue
        idx = remaining.lower().find(stripped.lower())
        if idx != -1:
            # Text trước match → thuộc run trước (nếu run i > 0 và run trước chưa có)
            if i > 0 and not distributed[i - 1] and idx > 0:
                distributed[i - 1] = remaining[:idx].rstrip()
            distributed[i] = remaining[idx: idx + len(stripped)]
            remaining = remaining[idx + len(stripped):].lstrip()
            matched_any = True

    if matched_any:
        # Phần còn lại của text dịch (sau tất cả các match) → gán vào run cuối có text
        if remaining:
            for i in range(len(runs) - 1, -1, -1):
                if distributed[i]:
                    distributed[i] = distributed[i] + " " + remaining.lstrip()
                    break
            else:
                distributed[-1] = remaining
        for i, run in enumerate(runs):
            _set_run_text(run, distributed[i])
        return

    # Bước 2 (fallback): phân phối tỉ lệ theo độ dài gốc, giữ ranh giới từ
    orig_total = sum(len(t) for t in orig_texts) or 1
    new_total = len(new_text)
    words = new_text.split()
    word_idx = 0
    for i, run in enumerate(runs):
        if i == len(runs) - 1:
            chunk = " ".join(words[word_idx:])
        else:
            ratio = len(orig_texts[i]) / orig_total
            target_len = max(1, round(new_total * ratio))
            chunk_words: list[str] = []
            chunk_len = 0
            while word_idx < len(words):
                w = words[word_idx]
                candidate = chunk_len + len(w) + (1 if chunk_words else 0)
                if candidate > target_len and chunk_words:
                    break
                chunk_words.append(w)
                chunk_len = candidate
                word_idx += 1
            chunk = " ".join(chunk_words)
        _set_run_text(run, chunk)


def _replace_cell_text_keep_fmt(cell: Any, new_text: str) -> None:
    """Thay text của table cell mà vẫn giữ formatting.
    Thao tác trên paragraph đầu tiên của cell; xóa các paragraph thừa.
    """
    try:
        from docx.oxml.ns import qn
    except ImportError:
        cell.text = new_text
        return

    paras = cell.paragraphs
    if not paras:
        cell.text = new_text
        return

    _replace_para_text_keep_fmt(paras[0], new_text)

    # Xóa bỏ các paragraph thừa (nếu cell gốc có nhiều paragraph)
    tc = cell._tc
    all_p = tc.findall(qn("w:p"))
    for extra_p in all_p[1:]:
        tc.remove(extra_p)


def _build_translated_docx(
    content: bytes,
    columns: List[str],
    rows: List[dict],
    to_translate: List[str],
) -> bytes:
    """Tạo lại file DOCX giữ nguyên cấu trúc và formatting.
    Luôn re-collect blocks theo thứ tự document (paragraphs + table cells)
    và áp dụng translation lên từng block tương ứng.
    """
    try:
        from docx import Document
    except ImportError:
        return content
    if not content or len(content) < 4:
        return content
    try:
        doc = Document(io.BytesIO(content))
    except Exception:
        return content

    blocks = _collect_docx_blocks(doc, max_items=len(rows))
    for i, row_data in enumerate(rows):
        if i >= len(blocks):
            break
        val = row_data.get("Nội dung_translated")
        if val is None:
            continue
        val = val.strip()
        _text, kind, ref = blocks[i]
        if kind == "para":
            _replace_para_text_keep_fmt(ref, val)
        else:  # kind == "cell"
            _replace_cell_text_keep_fmt(ref, val)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _read_full_excel(content: bytes, max_rows: int) -> Tuple[List[str], List[dict]]:
    """Đọc toàn bộ file .xlsx (openpyxl), trả về columns và toàn bộ rows (tối đa max_rows dòng dữ liệu).
    Bỏ qua cột trống trong header và hàng mà toàn bộ cell đều trống.
    """
    try:
        import openpyxl
    except ImportError:
        raise HTTPException(
            status_code=503,
            detail="openpyxl chưa được cài đặt. Vui lòng cài: pip install openpyxl",
        )
    workbook = openpyxl.load_workbook(io.BytesIO(content), read_only=True, data_only=True)
    sheet = workbook.active
    if not sheet:
        return [], []
    header_row = next(sheet.iter_rows(min_row=1, max_row=1, values_only=True), None)
    if not header_row:
        workbook.close()
        return [], []

    # Ghi nhận index các cột có tên (bỏ qua cột header trống)
    raw_columns = [_cell_to_str(h) for h in header_row]
    valid_col_indices = [i for i, c in enumerate(raw_columns) if c.strip()]

    seen: dict = {}
    columns: List[str] = []
    for i in valid_col_indices:
        c = raw_columns[i]
        key = c
        if key in seen:
            seen[key] += 1
            c = f"{c}_{seen[key]}"
        else:
            seen[key] = 1
        columns.append(c)

    rows = []
    counted = 0
    for row in sheet.iter_rows(min_row=2, values_only=True):
        if counted >= max_rows:
            break
        row_dict = {
            col_name: (_cell_to_str(row[orig_i]) if orig_i < len(row) else "")
            for col_name, orig_i in zip(columns, valid_col_indices)
        }
        # Bỏ qua hàng trống (toàn bộ cell đều rỗng)
        if not any(v.strip() for v in row_dict.values()):
            continue
        rows.append(row_dict)
        counted += 1
    workbook.close()
    return columns, rows


def _read_full_csv(content: bytes, max_rows: int) -> Tuple[List[str], List[dict]]:
    """Đọc toàn bộ file CSV, trả về columns và toàn bộ rows (tối đa max_rows dòng dữ liệu)."""
    try:
        text = content.decode("utf-8-sig").strip()
    except Exception:
        text = content.decode("utf-8", errors="replace").strip()
    reader = csv.reader(io.StringIO(text))
    header = next(reader, None)
    if not header:
        return [], []
    columns = [h.strip() or "Cột" for h in header]
    seen = {}
    unique_columns = []
    for c in columns:
        key = c
        if key in seen:
            seen[key] = seen.get(key, 1) + 1
            c = f"{c}_{seen[key]}"
        else:
            seen[key] = 1
        unique_columns.append(c)
    columns = unique_columns
    rows = []
    for _ in range(max_rows):
        row = next(reader, None)
        if row is None:
            break
        row_dict = {}
        for i, col_name in enumerate(columns):
            row_dict[col_name] = (row[i].strip() if i < len(row) else "")
        rows.append(row_dict)
    return columns, rows


@router.post("/parse-file", response_model=ParseFileResponse)
async def parse_file(file: UploadFile = File(...)):
    """
    Đọc file Excel/CSV/JSON/XML/DOCX, trả về danh sách cột và vài dòng xem trước.
    Dùng cho bước "Chọn cột" trên trang Dịch file.
    """
    if not file.filename:
        raise HTTPException(status_code=400, detail="Không có file.")
    name = (file.filename or "").lower()
    content = await file.read()
    if not content:
        raise HTTPException(status_code=400, detail="File rỗng.")
    if name.endswith(".xlsx"):
        return await _parse_excel(content)
    if name.endswith(".csv"):
        return await _parse_csv(content)
    if name.endswith(".json"):
        return await _parse_json(content)
    if name.endswith(".xml"):
        return await _parse_xml(content)
    if name.endswith(".docx"):
        return await _parse_docx(content)
    raise HTTPException(
        status_code=400,
        detail="Chỉ hỗ trợ file .xlsx, .csv, .json, .xml, .docx.",
    )


def _reconstruct_translated_json(
    content: bytes,
    rows: List[dict],
    to_translate: List[str],
) -> Optional[Dict[str, Any]]:
    """
    Tái tạo JSON gốc với các trường đã dịch, giữ nguyên cấu trúc nested.

    Xử lý 3 case:
      A. Root là list of objects      → ["users": [{"name": "..."}]]  → cập nhật từng item
      B. Root là dict chứa list       → {"users": [...]}               → tìm list, cập nhật
      C. Root là dict thuần (flat)    → {"title": "...", "ttl": 86400} → cập nhật key trực tiếp
         Nếu value là dict/list đã bị stringify → parse lại trước khi gán

    Khi value gốc là dict/list nhưng translated là JSON string → parse lại để giữ nested.
    """
    def _try_parse_json_value(translated_str: str, original_value: Any) -> Any:
        """Parse translated string về object nếu value gốc là dict/list."""
        if not isinstance(original_value, (dict, list)):
            return translated_str
        if not translated_str or not isinstance(translated_str, str):
            return original_value
        t = translated_str.strip()
        if not (t.startswith("{") or t.startswith("[")):
            return translated_str
        try:
            return json.loads(t)
        except (json.JSONDecodeError, ValueError):
            return translated_str

    def _merge_translated_row(target: dict, row: dict) -> None:
        """Gán giá trị đã dịch từ row vào target dict, parse JSON string nếu cần."""
        for col in to_translate:
            if col not in target:
                continue
            translated_val = row.get(col + "_translated")
            if translated_val is None:
                continue
            target[col] = _try_parse_json_value(translated_val, target[col])

    try:
        raw_text = _decode_text(content)
        data: Any = json.loads(raw_text)
    except (json.JSONDecodeError, TypeError, ValueError):
        return None

    # ── Case A: root là list of objects ────────────────────────────────────
    if isinstance(data, list) and len(data) > 0 and all(isinstance(x, dict) for x in data):
        for i in range(min(len(data), len(rows))):
            _merge_translated_row(data[i], rows[i])
        return data

    # ── Case B & C: root là dict ────────────────────────────────────────────
    if not isinstance(data, dict):
        return None

    # Case B: tìm key chứa list of objects
    for _k, v in data.items():
        if isinstance(v, list) and len(v) > 0 and all(isinstance(x, dict) for x in v):
            for i in range(min(len(v), len(rows))):
                _merge_translated_row(v[i], rows[i])
            return data

    # Case C: root dict thuần — các column là key của root, 1 row duy nhất
    if rows:
        _merge_translated_row(data, rows[0])
    return data


@router.post("/translate-file", response_model=TranslateFileResponse)
async def translate_file(
    file: UploadFile = File(...),
    selected_columns: str = Form(..., description="JSON array of column names to translate"),
    source_lang: str = Form(...),
    target_lang: str = Form(...),
    prompt_id: str = Form(""),
    context: str = Form(""),
    style: str = Form(""),
    game_id: str = Form(""),
    game_category_id: str = Form(""),
    db: AsyncSession = Depends(get_db),
):
    """
    Đọc file Excel/CSV, dịch các cột đã chọn theo từng ô (Cache -> Từ điển -> AI).
    Trả về toàn bộ dòng với cột gốc + cột _translated cho mỗi cột được chọn.
    Giới hạn tối đa TRANSLATE_FILE_MAX_ROWS dòng.
    """
    if not file.filename:
        raise HTTPException(status_code=400, detail="Không có file.")
    name = (file.filename or "").lower()
    content = await file.read()
    if not content:
        raise HTTPException(status_code=400, detail="File rỗng.")
    if not any(name.endswith(ext) for ext in (".xlsx", ".csv", ".json", ".xml", ".docx")):
        raise HTTPException(status_code=400, detail="Chỉ hỗ trợ file .xlsx, .csv, .json, .xml, .docx.")

    try:
        cols_list = json.loads(selected_columns)
        if not isinstance(cols_list, list):
            cols_list = []
    except (json.JSONDecodeError, TypeError):
        cols_list = []
    selected_set = {str(c).strip() for c in cols_list if c}

    source_lang = (source_lang or "").strip()
    target_lang = (target_lang or "").strip()
    if not source_lang or not target_lang:
        raise HTTPException(status_code=400, detail="Vui lòng chọn ngôn ngữ nguồn và ngôn ngữ đích.")
    if source_lang == target_lang:
        raise HTTPException(status_code=400, detail="Ngôn ngữ nguồn và đích phải khác nhau.")

    prompt_id_int = None
    if (prompt_id or "").strip():
        try:
            prompt_id_int = int(prompt_id.strip())
        except ValueError:
            pass
    game_id_int = None
    if (game_id or "").strip():
        try:
            game_id_int = int(game_id.strip())
        except ValueError:
            pass
    game_category_id_int = None
    if (game_category_id or "").strip():
        try:
            game_category_id_int = int(game_category_id.strip())
        except ValueError:
            pass
    context = (context or "").strip() or None
    style = (style or "").strip() or None

    if name.endswith(".xlsx"):
        columns, rows = _read_full_excel(content, TRANSLATE_FILE_MAX_ROWS)
    elif name.endswith(".csv"):
        columns, rows = _read_full_csv(content, TRANSLATE_FILE_MAX_ROWS)
    elif name.endswith(".json"):
        columns, rows = _read_full_json(content, TRANSLATE_FILE_MAX_ROWS)
    elif name.endswith(".xml"):
        columns, rows = _read_full_xml(content, TRANSLATE_FILE_MAX_ROWS)
    elif name.endswith(".docx"):
        columns, rows = _read_full_docx(content, TRANSLATE_FILE_MAX_ROWS)
    else:
        columns, rows = [], []

    to_translate = [c for c in selected_set if c in columns]
    if not to_translate:
        # Trả 200 với dữ liệu rỗng thay vì 400 (tránh lỗi "Failed to load resource" khi JSON/XML parse lỗi)
        fallback_cols = list(selected_set) if selected_set else (columns or ["Nội dung"])
        return TranslateFileResponse(columns=fallback_cols, rows=[])

    # ── Pass 1: tra Cache + Glossary (không gọi AI), ghi kết quả hoặc đánh dấu pending ──
    pending: List[Tuple[int, str, str]] = []
    for row_idx, row in enumerate(rows):
        for col in to_translate:
            text = (row.get(col) or "").strip()
            if not text:
                row[col + "_translated"] = ""
                continue
            try:
                hit = await translate_check_only(
                    db, text, source_lang, target_lang,
                    game_id=game_id_int, game_category_id=game_category_id_int,
                )
            except Exception as e:
                logger.warning("translate_check_only row=%s col=%s: %s", row_idx, col, e)
                try:
                    await db.rollback()
                except Exception:
                    pass
                hit = None
            if hit is not None:
                row[col + "_translated"] = hit
            else:
                row[col + "_translated"] = None
                pending.append((row_idx, col, text))

    # ── Pass 2: gom batch theo token → gọi AI 1 lần / batch → lưu cache ──
    token_batches = _build_token_batches(pending)
    for batch in token_batches:
        texts_to_translate = [t for _, _, t in batch]
        try:
            results = await translate_with_ai_batch(
                db,
                texts=texts_to_translate,
                source_lang=source_lang,
                target_lang=target_lang,
                prompt_id=prompt_id_int,
                context=context,
                style=style,
                game_id=game_id_int,
                game_category_id=game_category_id_int,
            )
        except Exception as e:
            logger.warning("translate_with_ai_batch: %s", e)
            results = [t for _, _, t in batch]
        saved_count = 0
        for (row_idx, col, orig_text), translated in zip(batch, results):
            rows[row_idx][col + "_translated"] = translated or orig_text
            if translated and translated != orig_text:
                await save_translation_to_cache(db, orig_text, translated, source_lang, target_lang, origin="file")
                saved_count += 1
            logger.info("translate_file batch: %d results, %d cached", len(batch), saved_count)

    output_columns = columns + [c + "_translated" for c in to_translate]
    translated_json = None
    translated_docx_b64 = None
    if name.endswith(".json"):
        translated_json = _reconstruct_translated_json(content, rows, to_translate)
    elif name.endswith(".docx"):
        try:
            docx_bytes = _build_translated_docx(content, columns, rows, to_translate)
            translated_docx_b64 = base64.b64encode(docx_bytes).decode("ascii")
        except Exception as e:
            logger.warning("Build translated DOCX: %s", e)
    return TranslateFileResponse(
        columns=output_columns,
        rows=rows,
        translated_json=translated_json,
        translated_docx_b64=translated_docx_b64,
    )


@router.post("/translate-file-stream")
async def translate_file_stream(
    file: UploadFile = File(...),
    selected_columns: str = Form(..., description="JSON array of column names to translate"),
    source_lang: str = Form(...),
    target_lang: str = Form(...),
    prompt_id: str = Form(""),
    context: str = Form(""),
    style: str = Form(""),
    game_id: str = Form(""),
    game_category_id: str = Form(""),
    db: AsyncSession = Depends(get_db),
):
    """
    Dịch file với Server-Sent Events (SSE) để hiển thị tiến trình real-time.
    Trả về stream text/event-stream với các event: start, progress, done, error.
    """
    from fastapi.responses import StreamingResponse as _SR

    if not file.filename:
        raise HTTPException(status_code=400, detail="Không có file.")
    name = (file.filename or "").lower()
    content = await file.read()
    if not content:
        raise HTTPException(status_code=400, detail="File rỗng.")
    if not any(name.endswith(ext) for ext in (".xlsx", ".csv", ".json", ".xml", ".docx")):
        raise HTTPException(status_code=400, detail="Chỉ hỗ trợ file .xlsx, .csv, .json, .xml, .docx.")

    try:
        cols_list = json.loads(selected_columns)
        if not isinstance(cols_list, list):
            cols_list = []
    except (json.JSONDecodeError, TypeError):
        cols_list = []
    selected_set = {str(c).strip() for c in cols_list if c}
    source_lang = (source_lang or "").strip()
    target_lang = (target_lang or "").strip()
    if not source_lang or not target_lang:
        raise HTTPException(status_code=400, detail="Vui lòng chọn ngôn ngữ nguồn và ngôn ngữ đích.")
    if source_lang == target_lang:
        raise HTTPException(status_code=400, detail="Ngôn ngữ nguồn và đích phải khác nhau.")

    prompt_id_int: Optional[int] = None
    if (prompt_id or "").strip():
        try:
            prompt_id_int = int(prompt_id.strip())
        except ValueError:
            pass
    game_id_int: Optional[int] = None
    if (game_id or "").strip():
        try:
            game_id_int = int(game_id.strip())
        except ValueError:
            pass
    game_category_id_int: Optional[int] = None
    if (game_category_id or "").strip():
        try:
            game_category_id_int = int(game_category_id.strip())
        except ValueError:
            pass
    ctx = (context or "").strip() or None
    sty = (style or "").strip() or None

    # Đọc file
    if name.endswith(".xlsx"):
        columns, rows = _read_full_excel(content, TRANSLATE_FILE_MAX_ROWS)
    elif name.endswith(".csv"):
        columns, rows = _read_full_csv(content, TRANSLATE_FILE_MAX_ROWS)
    elif name.endswith(".json"):
        columns, rows = _read_full_json(content, TRANSLATE_FILE_MAX_ROWS)
    elif name.endswith(".xml"):
        columns, rows = _read_full_xml(content, TRANSLATE_FILE_MAX_ROWS)
    elif name.endswith(".docx"):
        columns, rows = _read_full_docx(content, TRANSLATE_FILE_MAX_ROWS)
    else:
        columns, rows = [], []

    to_translate = [c for c in selected_set if c in columns]

    async def generate():
        def _sse(obj: dict) -> str:
            return f"data: {json.dumps(obj, ensure_ascii=False)}\n\n"

        if not to_translate:
            yield _sse({"type": "done", "columns": list(selected_set) or columns or ["Nội dung"], "rows": [], "translated_json": None, "translated_docx_b64": None})
            return

        # ── Pass 1: Cache + Glossary ──
        pending: List[Tuple[int, str, str]] = []
        for row_idx, row in enumerate(rows):
            for col in to_translate:
                text = (row.get(col) or "").strip()
                if not text:
                    row[col + "_translated"] = ""
                    continue
                try:
                    hit = await translate_check_only(
                        db, text, source_lang, target_lang,
                        game_id=game_id_int, game_category_id=game_category_id_int,
                    )
                except Exception as e:
                    logger.warning("SSE translate_check_only row=%s col=%s: %s", row_idx, col, e)
                    try:
                        await db.rollback()
                    except Exception:
                        pass
                    hit = None
                if hit is not None:
                    row[col + "_translated"] = hit
                else:
                    row[col + "_translated"] = None
                    pending.append((row_idx, col, text))

        total_cells = len(pending)
        token_batches = _build_token_batches(pending)
        batch_total = len(token_batches)

        yield _sse({"type": "start", "total": total_cells, "batch_total": batch_total})

        logger.info("SSE file translate: %d pending cells, %d batches", total_cells, batch_total)
        cells_done = 0
        for batch_idx, batch in enumerate(token_batches):
            texts_to_translate = [t for _, _, t in batch]
            logger.info("SSE batch %s input (first 3): %s", batch_idx, texts_to_translate[:3])
            try:
                results = await translate_with_ai_batch(
                    db,
                    texts=texts_to_translate,
                    source_lang=source_lang,
                    target_lang=target_lang,
                    prompt_id=prompt_id_int,
                    context=ctx,
                    style=sty,
                    game_id=game_id_int,
                    game_category_id=game_category_id_int,
                )
                logger.info("SSE batch %s output (first 3): %s", batch_idx, results[:3])
            except Exception as e:
                logger.warning("SSE translate batch %s FAILED: %s", batch_idx, e, exc_info=True)
                results = [t for _, _, t in batch]

            saved_count = 0
            for (row_idx, col, orig_text), translated in zip(batch, results):
                rows[row_idx][col + "_translated"] = translated or orig_text
                if translated and translated != orig_text:
                    await save_translation_to_cache(db, orig_text, translated, source_lang, target_lang, origin="file")
                    saved_count += 1
            logger.info("SSE batch %s: %d results, %d cached", batch_idx, len(batch), saved_count)

            cells_done += len(batch)
            percent = round(cells_done * 100 / total_cells) if total_cells else 100
            yield _sse({
                "type": "progress",
                "done": cells_done,
                "total": total_cells,
                "batch_done": batch_idx + 1,
                "batch_total": batch_total,
                "percent": percent,
                "batch_size": len(batch),
                "batch_tokens": sum(_estimate_tokens(t) for _, _, t in batch),
            })

        # Xây dựng kết quả cuối
        output_columns = columns + [c + "_translated" for c in to_translate]
        translated_json = None
        translated_docx_b64 = None
        if name.endswith(".json"):
            translated_json = _reconstruct_translated_json(content, rows, to_translate)
        elif name.endswith(".docx"):
            try:
                docx_bytes = _build_translated_docx(content, columns, rows, to_translate)
                translated_docx_b64 = base64.b64encode(docx_bytes).decode("ascii")
            except Exception as e:
                logger.warning("SSE build DOCX: %s", e)

        yield _sse({
            "type": "done",
            "columns": output_columns,
            "rows": rows,
            "translated_json": translated_json,
            "translated_docx_b64": translated_docx_b64,
        })

    return _SR(
        generate(),
        media_type="text/event-stream",
        headers={
            "Cache-Control": "no-cache",
            "X-Accel-Buffering": "no",
            "Connection": "keep-alive",
        },
    )


@router.post("/translate-json-file")
async def translate_json_file(
    file: UploadFile = File(...),
    source_lang: str = Form(...),
    target_lang: str = Form(...),
    smart_filter: str = Form("true"),
    translate_keys: str = Form("false"),
    prompt_id: str = Form(""),
    context: str = Form(""),
    style: str = Form(""),
    game_id: str = Form(""),
    game_category_id: str = Form(""),
    db: AsyncSession = Depends(get_db),
):
    """
    Dịch file JSON giữ nguyên cấu trúc: chỉ dịch value là chuỗi (Cache → Từ điển → AI).
    Không đụng key, số, boolean, null, cấu trúc. Smart Filter: bỏ qua URL, email, số, ngày ISO, enum.
    Trả về JSON (application/json), frontend có thể lưu thành file.
    """
    if not file.filename or not (file.filename or "").lower().endswith(".json"):
        raise HTTPException(status_code=400, detail="Chỉ chấp nhận file .json.")
    content = await file.read()
    if not content:
        raise HTTPException(status_code=400, detail="File rỗng.")
    text = _decode_text(content)
    try:
        data = json.loads(text)
    except json.JSONDecodeError:
        raise HTTPException(status_code=400, detail="File không phải JSON hợp lệ. Vui lòng kiểm tra cú pháp.")
    source_lang = (source_lang or "").strip()
    target_lang = (target_lang or "").strip()
    if not source_lang or not target_lang:
        raise HTTPException(status_code=400, detail="Vui lòng chọn ngôn ngữ nguồn và ngôn ngữ đích.")
    if source_lang == target_lang:
        raise HTTPException(status_code=400, detail="Ngôn ngữ nguồn và đích phải khác nhau.")
    use_smart_filter = (smart_filter or "").strip().lower() in ("true", "1", "yes")
    use_translate_keys = (translate_keys or "").strip().lower() in ("true", "1", "yes")
    prompt_id_int = None
    if (prompt_id or "").strip():
        try:
            prompt_id_int = int(prompt_id.strip())
        except ValueError:
            pass
    game_id_int = None
    if (game_id or "").strip():
        try:
            game_id_int = int(game_id.strip())
        except ValueError:
            pass
    game_category_id_int = None
    if (game_category_id or "").strip():
        try:
            game_category_id_int = int(game_category_id.strip())
        except ValueError:
            pass
    context = (context or "").strip() or None
    style = (style or "").strip() or None

    result = await _translate_json_recursive(
        data,
        db,
        source_lang,
        target_lang,
        prompt_id_int,
        context,
        style,
        smart_filter=use_smart_filter,
        translate_keys=use_translate_keys,
        game_id=game_id_int,
        game_category_id=game_category_id_int,
    )
    body = json.dumps(result, ensure_ascii=False, indent=2)
    filename = (file.filename or "translated").replace(".json", "") + "_translated.json"
    return Response(
        content=body.encode("utf-8"),
        media_type="application/json",
        headers={"Content-Disposition": f'attachment; filename="{filename}"'},
    )


@router.post("/translate-xml-file")
async def translate_xml_file(
    file: UploadFile = File(...),
    source_lang: str = Form(...),
    target_lang: str = Form(...),
    preserve_placeholders: str = Form("true"),
    respect_translatable: str = Form("true"),
    smart_filter: str = Form("true"),
    prompt_id: str = Form(""),
    context: str = Form(""),
    style: str = Form(""),
    game_id: str = Form(""),
    game_category_id: str = Form(""),
    db: AsyncSession = Depends(get_db),
):
    """
    Dịch file XML giữ cấu trúc: chỉ dịch text bên trong thẻ (string, text, description, title...).
    Không dịch attribute, tên thẻ. Giữ placeholder {username}, %d, %s, {{count}}.
    Chỉ dịch thẻ translatable="true" hoặc thẻ UI; bỏ qua translatable="false".
    """
    if not file.filename or not (file.filename or "").lower().endswith(".xml"):
        raise HTTPException(status_code=400, detail="Chỉ chấp nhận file .xml.")
    content = await file.read()
    if not content:
        raise HTTPException(status_code=400, detail="File rỗng.")
    text = _decode_text(content)
    decl, rest = _split_xml_declaration(text)
    if not rest.strip():
        raise HTTPException(status_code=400, detail="File XML không có nội dung.")
    try:
        root = ET.fromstring(rest)
    except ET.ParseError as e:
        raise HTTPException(status_code=400, detail=f"File không phải XML hợp lệ: {str(e)[:150]}")
    source_lang = (source_lang or "").strip()
    target_lang = (target_lang or "").strip()
    if not source_lang or not target_lang:
        raise HTTPException(status_code=400, detail="Vui lòng chọn ngôn ngữ nguồn và ngôn ngữ đích.")
    if source_lang == target_lang:
        raise HTTPException(status_code=400, detail="Ngôn ngữ nguồn và đích phải khác nhau.")
    use_preserve_placeholders = (preserve_placeholders or "").strip().lower() in ("true", "1", "yes")
    use_respect_translatable = (respect_translatable or "").strip().lower() in ("true", "1", "yes")
    use_smart_filter = (smart_filter or "").strip().lower() in ("true", "1", "yes")
    prompt_id_int = None
    if (prompt_id or "").strip():
        try:
            prompt_id_int = int(prompt_id.strip())
        except ValueError:
            pass
    game_id_int = None
    if (game_id or "").strip():
        try:
            game_id_int = int(game_id.strip())
        except ValueError:
            pass
    game_category_id_int = None
    if (game_category_id or "").strip():
        try:
            game_category_id_int = int(game_category_id.strip())
        except ValueError:
            pass
    context = (context or "").strip() or None
    style = (style or "").strip() or None

    await _translate_xml_recursive(
        root,
        db,
        source_lang,
        target_lang,
        prompt_id_int,
        context,
        style,
        preserve_placeholders=use_preserve_placeholders,
        respect_translatable=use_respect_translatable,
        smart_filter=use_smart_filter,
        game_id=game_id_int,
        game_category_id=game_category_id_int,
    )
    out_body = (_xml_to_string(root, declaration=decl) if decl else _xml_to_string(root))
    filename = (file.filename or "translated").replace(".xml", "") + "_translated.xml"
    return Response(
        content=out_body.encode("utf-8"),
        media_type="application/xml",
        headers={"Content-Disposition": f'attachment; filename="{filename}"'},
    )


def _export_csv(columns: List[str], rows: List[dict]) -> bytes:
    out = io.StringIO()
    out.write("\uFEFF")
    writer = csv.writer(out, lineterminator="\r\n")
    writer.writerow(columns)
    for row in rows:
        writer.writerow([str(row.get(c, "") or "") for c in columns])
    return out.getvalue().encode("utf-8")


def _export_xlsx(columns: List[str], rows: List[dict]) -> bytes:
    try:
        import openpyxl
    except ImportError:
        raise HTTPException(status_code=503, detail="openpyxl chưa được cài đặt.")
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.append(columns)
    for row in rows:
        ws.append([str(row.get(c, "") or "") for c in columns])
    buf = io.BytesIO()
    wb.save(buf)
    return buf.getvalue()


def _export_json(columns: List[str], rows: List[dict]) -> bytes:
    data = [dict((c, row.get(c, "") or "") for c in columns) for row in rows]
    return json.dumps(data, ensure_ascii=False, indent=2).encode("utf-8")


def _xml_tag(s: str) -> str:
    """Tên thẻ XML hợp lệ (chữ, số, _)."""
    return "".join(c if c.isalnum() or c == "_" else "_" for c in (s or "col"))


def _export_xml(columns: List[str], rows: List[dict]) -> bytes:
    lines = ['<?xml version="1.0" encoding="UTF-8"?>', "<rows>"]
    for row in rows:
        lines.append("  <row>")
        for c in columns:
            val = (row.get(c) or "").replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;").replace('"', "&quot;")
            tag = _xml_tag(c)
            lines.append(f"    <{tag}>{val}</{tag}>")
        lines.append("  </row>")
    lines.append("</rows>")
    return "\n".join(lines).encode("utf-8")


def _export_docx(columns: List[str], rows: List[dict]) -> bytes:
    try:
        from docx import Document
    except ImportError:
        raise HTTPException(status_code=503, detail="python-docx chưa được cài đặt.")
    doc = Document()
    table = doc.add_table(rows=1 + len(rows), cols=len(columns))
    for j, col in enumerate(columns):
        table.rows[0].cells[j].text = str(col)
    for i, row in enumerate(rows):
        for j, col in enumerate(columns):
            table.rows[i + 1].cells[j].text = str(row.get(col, "") or "")
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


@router.post("/parse-file-full", response_model=ParseAllRowsResponse)
async def parse_file_full(file: UploadFile = File(...)):
    """
    Parse toàn bộ dòng từ file .xlsx hoặc .csv (không giới hạn 5 dòng preview).
    Dùng cho tính năng Hiệu Đính File - cần đọc tất cả dòng để chỉnh sửa/hiệu đính.
    Giới hạn tối đa TRANSLATE_FILE_MAX_ROWS dòng.
    """
    if not file.filename:
        raise HTTPException(status_code=400, detail="Không có file.")
    filename_lower = (file.filename or "").lower()
    content = await file.read()
    if not content:
        raise HTTPException(status_code=400, detail="File rỗng.")

    if filename_lower.endswith(".xlsx") or filename_lower.endswith(".xls"):
        try:
            import openpyxl
        except ImportError:
            raise HTTPException(status_code=503, detail="openpyxl chưa được cài đặt.")
        workbook = openpyxl.load_workbook(io.BytesIO(content), read_only=True, data_only=True)
        sheet = workbook.active
        if not sheet:
            return ParseAllRowsResponse(columns=[], rows=[], total=0)
        header_row = next(sheet.iter_rows(min_row=1, max_row=1, values_only=True), None)
        if not header_row:
            return ParseAllRowsResponse(columns=[], rows=[], total=0)
        raw_columns = [_cell_to_str(h) for h in header_row]
        valid_col_indices = [i for i, c in enumerate(raw_columns) if c.strip()]
        seen: dict = {}
        columns: List[str] = []
        for i in valid_col_indices:
            c = raw_columns[i]
            if c in seen:
                seen[c] += 1
                c = f"{c}_{seen[c]}"
            else:
                seen[c] = 1
            columns.append(c)
        rows: List[dict] = []
        for row in sheet.iter_rows(min_row=2, max_row=1 + TRANSLATE_FILE_MAX_ROWS, values_only=True):
            row_dict = {
                col_name: (_cell_to_str(row[orig_i]) if orig_i < len(row) else "")
                for col_name, orig_i in zip(columns, valid_col_indices)
            }
            if not any(v.strip() for v in row_dict.values()):
                continue
            rows.append(row_dict)
        workbook.close()
        return ParseAllRowsResponse(columns=columns, rows=rows, total=len(rows))

    elif filename_lower.endswith(".csv"):
        try:
            text = content.decode("utf-8-sig").strip()
        except Exception:
            text = content.decode("utf-8", errors="replace").strip()
        reader = csv.reader(io.StringIO(text))
        header = next(reader, None)
        if not header:
            return ParseAllRowsResponse(columns=[], rows=[], total=0)
        raw_cols = [h.strip() or "Cột" for h in header]
        seen2: dict = {}
        columns2: List[str] = []
        for c in raw_cols:
            if c in seen2:
                seen2[c] += 1
                c = f"{c}_{seen2[c]}"
            else:
                seen2[c] = 1
            columns2.append(c)
        rows2: List[dict] = []
        for i, row in enumerate(reader):
            if i >= TRANSLATE_FILE_MAX_ROWS:
                break
            row_dict = {col: (row[j].strip() if j < len(row) else "") for j, col in enumerate(columns2)}
            rows2.append(row_dict)
        return ParseAllRowsResponse(columns=columns2, rows=rows2, total=len(rows2))

    else:
        raise HTTPException(status_code=400, detail="Chỉ chấp nhận file .xlsx hoặc .csv cho Hiệu Đính File.")


@router.post("/proofread-row", response_model=ProofreadRowResponse)
async def proofread_row_endpoint(
    body: ProofreadRowRequest,
    db: AsyncSession = Depends(get_db),
):
    """
    Hiệu đính 1 dòng bằng AI: nhận văn bản gốc + bản dịch hiện tại, trả về bản dịch đã cải thiện.
    """
    if not body.original.strip():
        raise HTTPException(status_code=400, detail="Văn bản gốc không được để trống.")
    if not body.source_lang.strip() or not body.target_lang.strip():
        raise HTTPException(status_code=400, detail="Vui lòng chọn ngôn ngữ nguồn và đích.")
    try:
        result = await proofread_with_ai(
            db,
            original=body.original.strip(),
            translated=(body.translated or "").strip(),
            source_lang=body.source_lang.strip(),
            target_lang=body.target_lang.strip(),
            prompt_id=body.prompt_id,
            context=body.context or None,
            style=body.style or None,
        )
    except ValueError as e:
        raise HTTPException(status_code=400, detail=str(e))
    except Exception as e:
        raise HTTPException(status_code=502, detail=f"Hiệu đính AI thất bại: {getattr(e, 'message', str(e))}")
    if result:
        try:
            await save_translation_to_cache(
                db, body.original.strip(), result,
                body.source_lang.strip(), body.target_lang.strip(),
                origin="proofread",
            )
        except Exception:
            pass
    return ProofreadRowResponse(proofread=result)


@router.post("/proofread-batch", response_model=ProofreadBatchResponse)
async def proofread_batch_endpoint(
    body: ProofreadBatchRequest,
    db: AsyncSession = Depends(get_db),
):
    """
    Hiệu đính nhiều dòng cùng lúc (batch AI). Mỗi item gồm index, original, translated.
    Trả về danh sách kết quả [{index, proofread}].
    """
    if not body.items:
        raise HTTPException(status_code=400, detail="Danh sách dòng không được rỗng.")
    if not body.source_lang.strip() or not body.target_lang.strip():
        raise HTTPException(status_code=400, detail="Vui lòng chọn ngôn ngữ nguồn và đích.")
    items_dicts = [
        {"index": item.index, "original": item.original, "translated": item.translated}
        for item in body.items
    ]
    try:
        results = await proofread_with_ai_batch(
            db,
            items=items_dicts,
            source_lang=body.source_lang.strip(),
            target_lang=body.target_lang.strip(),
            prompt_id=body.prompt_id,
            context=body.context or None,
            style=body.style or None,
        )
    except ValueError as e:
        raise HTTPException(status_code=400, detail=str(e))
    except Exception as e:
        raise HTTPException(status_code=502, detail=f"Hiệu đính batch AI thất bại: {getattr(e, 'message', str(e))}")
    items_by_index = {item["index"]: item for item in items_dicts}
    for r in results:
        orig_item = items_by_index.get(r["index"])
        if orig_item and r.get("proofread"):
            try:
                await save_translation_to_cache(
                    db, (orig_item.get("original") or "").strip(), r["proofread"],
                    body.source_lang.strip(), body.target_lang.strip(),
                    origin="proofread",
                )
            except Exception:
                pass
    return ProofreadBatchResponse(
        results=[ProofreadBatchResultItem(index=r["index"], proofread=r["proofread"]) for r in results]
    )


@router.post("/export-file")
async def export_file(body: ExportFileRequest):
    """
    Xuất dữ liệu đã dịch (columns + rows) ra file theo định dạng: csv, xlsx, json, xml, docx.
    Frontend gọi với format trùng đuôi file đã upload để tải về đúng định dạng.
    """
    fmt = (body.format or "").lower().strip()
    if fmt not in ("csv", "xlsx", "json", "xml", "docx"):
        raise HTTPException(status_code=400, detail="format phải là một trong: csv, xlsx, json, xml, docx.")
    columns = body.columns or []
    rows = body.rows or []
    if not columns:
        raise HTTPException(status_code=400, detail="columns không được rỗng.")
    if fmt == "csv":
        content = _export_csv(columns, rows)
        media_type = "text/csv; charset=utf-8"
    elif fmt == "xlsx":
        content = _export_xlsx(columns, rows)
        media_type = "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
    elif fmt == "json":
        content = _export_json(columns, rows)
        media_type = "application/json; charset=utf-8"
    elif fmt == "xml":
        content = _export_xml(columns, rows)
        media_type = "application/xml; charset=utf-8"
    else:
        content = _export_docx(columns, rows)
        media_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    base = (body.filename or "dich").strip() or "dich"
    filename = f"{base}_translated.{fmt}"
    return Response(
        content=content,
        media_type=media_type,
        headers={"Content-Disposition": f'attachment; filename="{filename}"'},
    )


@router.get("/verify")
async def verify_api_key(db: AsyncSession = Depends(get_db)):
    """
    Kiểm tra Gemini API key có hoạt động hay không (đọc từ Cài đặt).
    Luôn trả 200 với { "ok": true/false, "message": "..." } để client không bị "Failed to load resource".
    """
    ok, message = await verify_gemini_api_key(db)
    return {"ok": ok, "message": message}


@router.get("/test-batch-cache")
async def test_batch_cache(
    source_lang: str = Query("vi", description="Ngôn ngữ nguồn"),
    target_lang: str = Query("en", description="Ngôn ngữ đích"),
    db: AsyncSession = Depends(get_db),
):
    """
    Test end-to-end: batch translate 2 dòng test + save to cache.
    Kiểm tra xem AI batch và cache saving có hoạt động không.
    """
    test_texts = ["Xin chào", "Cảm ơn"]
    result: dict = {
        "input": test_texts,
        "translated": [],
        "cached_count": 0,
        "ai_error": None,
        "cache_error": None,
    }
    try:
        from app.modules.translate.service import translate_with_ai_batch, save_translation_to_cache
        translated = await translate_with_ai_batch(
            db,
            texts=test_texts,
            source_lang=source_lang,
            target_lang=target_lang,
        )
        result["translated"] = translated
    except Exception as e:
        result["ai_error"] = str(e)
        return result

    for orig, trans in zip(test_texts, translated):
        if trans and trans != orig:
            try:
                await save_translation_to_cache(db, orig, trans, source_lang, target_lang, origin="direct")
                result["cached_count"] += 1
            except Exception as e:
                result["cache_error"] = str(e)

    return result


@router.get("/cache-diagnostics")
async def cache_diagnostics(db: AsyncSession = Depends(get_db)):
    """
    Endpoint chẩn đoán cache: ghi 1 entry test, đọc lại, rồi xóa.
    Trả về { "write_ok", "read_ok", "error", "cache_count" }.
    """
    from app.modules.cache.service import CacheService as _CS
    svc = _CS(db)
    test_key = "translate:test:test:diag0000deadbeef"
    result: dict = {"write_ok": False, "read_ok": False, "error": None, "cache_count": 0}
    try:
        # Đếm tổng số entries hiện tại
        listed = await svc.list(skip=0, limit=1)
        result["cache_count"] = listed.get("total", 0)
        # Thử ghi
        await svc.create({"key": test_key, "value": "__diag_test__", "ttl": 3600})
        result["write_ok"] = True
        # Thử đọc lại
        entry = await svc.get_by_key(test_key)
        result["read_ok"] = entry is not None and getattr(entry, "value", "") == "__diag_test__"
        # Xóa entry test
        if entry:
            await svc.delete(entry.id)
    except Exception as e:
        result["error"] = str(e)
    return result


@router.post("", response_model=TranslateResponse)
async def translate(
    body: TranslateRequest,
    db: AsyncSession = Depends(get_db),
):
    """
    Dịch văn bản theo thứ tự ưu tiên: Cache -> Từ điển game -> Từ điển chung -> AI (Google Gemini).
    Nếu dùng AI, kết quả được lưu cache cho lần sau. Prompt: prompt_id từ Quản lý Prompts hoặc mặc định.
    """
    if not body.text or not body.text.strip():
        logger.info("Translate 400: empty text")
        raise HTTPException(status_code=400, detail="Văn bản cần dịch không được để trống.")
    if not (body.source_lang or "").strip() or not (body.target_lang or "").strip():
        logger.info("Translate 400: missing source or target language")
        raise HTTPException(status_code=400, detail="Vui lòng chọn ngôn ngữ nguồn và ngôn ngữ đích.")
    if body.source_lang == body.target_lang:
        logger.info("Translate 400: source_lang == target_lang")
        raise HTTPException(status_code=400, detail="Ngôn ngữ nguồn và đích phải khác nhau.")

    try:
        translated_text = await translate_with_priority(
            db,
            text=body.text.strip(),
            source_lang=body.source_lang.strip(),
            target_lang=body.target_lang.strip(),
            prompt_id=body.prompt_id,
            context=body.context or None,
            style=body.style or None,
        )
    except ValueError as e:
        msg = str(e)
        logger.warning("Translate error (ValueError): %s", msg)
        raise HTTPException(status_code=400, detail=msg)
    except Exception as e:
        logger.warning("Translate AI failed, returning original text: %s", e)
        return TranslateResponse(translated_text=body.text.strip())

    return TranslateResponse(translated_text=translated_text)
