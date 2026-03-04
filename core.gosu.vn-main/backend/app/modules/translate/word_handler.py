"""
Xử lý file Word (.docx): đọc, parse preview, đọc full, build file đã dịch, xuất.
"""
import io
import re as _re
from typing import Any, List, Tuple

from fastapi import HTTPException

from app.modules.translate.schemas import ParseFileResponse


def _iter_docx_body_items(doc: Any):
    """Yield Paragraph hoặc Table theo thứ tự xuất hiện trong document body."""
    try:
        from docx.text.paragraph import Paragraph as _Para
        from docx.table import Table as _Tbl
    except ImportError:
        return
    for child in doc.element.body.iterchildren():
        tag = child.tag.split("}")[-1] if "}" in child.tag else child.tag
        if tag == "p":
            yield _Para(child, doc)
        elif tag == "tbl":
            yield _Tbl(child, doc)


def _collect_docx_blocks(doc: Any, max_items: int = 100_000) -> List[Tuple[str, str, Any]]:
    """Thu thập text blocks theo thứ tự: paragraph + table cells. Returns (text, kind, ref)."""
    try:
        from docx.text.paragraph import Paragraph as _Para
        from docx.table import Table as _Tbl
    except ImportError:
        return []
    blocks: List[Tuple[str, str, Any]] = []
    for item in _iter_docx_body_items(doc):
        if len(blocks) >= max_items:
            break
        if isinstance(item, _Para):
            text = (item.text or "").strip()
            if text:
                blocks.append((text, "para", item))
        elif isinstance(item, _Tbl):
            seen: set = set()
            for row in item.rows:
                for cell in row.cells:
                    cid = id(cell._tc)
                    if cid in seen:
                        continue
                    seen.add(cid)
                    text = (cell.text or "").strip()
                    if text:
                        blocks.append((text, "cell", cell))
    return blocks


def _para_to_html(para: Any) -> str:
    """Chuyển 1 paragraph python-docx thành HTML, giữ heading/bold/italic/list/underline."""
    import html as _html_mod
    style_name: str = (para.style.name or "") if para.style else ""
    inner = ""
    for run in para.runs:
        text = run.text or ""
        if not text:
            continue
        text = _html_mod.escape(text)
        if run.bold and run.italic:
            text = f"<strong><em>{text}</em></strong>"
        elif run.bold:
            text = f"<strong>{text}</strong>"
        elif run.italic:
            text = f"<em>{text}</em>"
        if getattr(run, "underline", False):
            text = f"<u>{text}</u>"
        inner += text
    if not inner.strip():
        return ""
    sn_lower = style_name.lower()
    if "heading 1" in sn_lower:
        return f"<h1>{inner}</h1>"
    if "heading 2" in sn_lower:
        return f"<h2>{inner}</h2>"
    if "heading 3" in sn_lower or "heading 4" in sn_lower:
        return f"<h3>{inner}</h3>"
    if "list" in sn_lower or sn_lower.startswith("list"):
        return f"<li>{inner}</li>"
    return f"<p>{inner}</p>"


def _table_to_html(table: Any) -> str:
    """Chuyển 1 python-docx Table thành HTML <table> với border."""
    import html as _html_mod
    rows_html = ""
    for row_idx, row in enumerate(table.rows):
        seen: set = set()
        cells_html = ""
        for cell in row.cells:
            cid = id(cell._tc)
            if cid in seen:
                continue
            seen.add(cid)
            cell_inner = ""
            for p in cell.paragraphs:
                ph = _para_to_html(p)
                if ph:
                    m = _re.match(r"^<[^>]+>(.*)</[^>]+>$", ph, _re.DOTALL)
                    cell_inner += (m.group(1) if m else _html_mod.escape(cell.text)) + " "
            cell_inner = cell_inner.strip() or _html_mod.escape(cell.text or "")
            tag = "th" if row_idx == 0 else "td"
            cells_html += f"<{tag} style='border:1px solid #d1d5db;padding:6px 10px;text-align:left'>{cell_inner}</{tag}>"
        rows_html += f"<tr>{cells_html}</tr>"
    return (
        "<div style='overflow-x:auto;margin:8px 0'>"
        "<table style='border-collapse:collapse;width:100%;font-size:13px'>"
        f"{rows_html}"
        "</table></div>"
    )


def _docx_to_preview_html(doc: Any, max_items: int = 120) -> str:
    """Chuyển document (paragraphs + tables) thành HTML theo thứ tự body."""
    try:
        from docx.text.paragraph import Paragraph as _Para
        from docx.table import Table as _Tbl
    except ImportError:
        parts = []
        for p in list(doc.paragraphs)[:max_items]:
            h = _para_to_html(p)
            if h:
                parts.append(h)
        return "".join(parts)
    parts: list = []
    count = 0
    for item in _iter_docx_body_items(doc):
        if count >= max_items:
            break
        if isinstance(item, _Para):
            h = _para_to_html(item)
            if h:
                parts.append(h)
                count += 1
        elif isinstance(item, _Tbl):
            parts.append(_table_to_html(item))
            count += 1
    result = ""
    i = 0
    while i < len(parts):
        if parts[i].startswith("<li>"):
            result += "<ul>"
            while i < len(parts) and parts[i].startswith("<li>"):
                result += parts[i]
                i += 1
            result += "</ul>"
        else:
            result += parts[i]
            i += 1
    return result


async def parse_docx(content: bytes, preview_limit: int = 5) -> ParseFileResponse:
    """Đọc file .docx, trả về columns và preview_rows (preview_html cho UI)."""
    try:
        from docx import Document
    except ImportError:
        raise HTTPException(
            status_code=503,
            detail="Hỗ trợ DOCX cần cài thư viện: pip install python-docx (sau đó khởi động lại backend).",
        )
    if not content or len(content) < 4:
        raise HTTPException(status_code=400, detail="File DOCX rỗng hoặc không đúng định dạng.")
    try:
        doc = Document(io.BytesIO(content))
    except Exception as e:
        raise HTTPException(status_code=400, detail=f"Không đọc được file DOCX (file có thể hỏng hoặc không phải .docx): {str(e)[:200]}")

    blocks = _collect_docx_blocks(doc, max_items=preview_limit)
    columns = ["Nội dung"]
    preview_rows = [{"Nội dung": text} for text, _kind, _ref in blocks]
    preview_html = _docx_to_preview_html(doc)
    return ParseFileResponse(columns=columns, preview_rows=preview_rows, preview_html=preview_html)


def read_full_docx(content: bytes, max_rows: int) -> Tuple[List[str], List[dict]]:
    """Đọc toàn bộ file .docx, trả về columns và rows (tối đa max_rows)."""
    try:
        from docx import Document
    except ImportError:
        raise HTTPException(
            status_code=503,
            detail="Hỗ trợ DOCX cần cài: pip install python-docx và khởi động lại backend.",
        )
    if not content or len(content) < 4:
        return [], []
    try:
        doc = Document(io.BytesIO(content))
    except Exception:
        return [], []
    blocks = _collect_docx_blocks(doc, max_items=max_rows)
    rows = [{"Nội dung": text} for text, _kind, _ref in blocks]
    return ["Nội dung"], rows


def _replace_para_text_keep_fmt(para: Any, new_text: str) -> None:
    """Thay text của paragraph, giữ formatting (bold/italic/underline...)."""
    try:
        from docx.oxml.ns import qn
        from lxml import etree
    except ImportError:
        para.text = new_text
        return

    XML_SPACE = "{http://www.w3.org/XML/1998/namespace}space"

    def _set_run_text(r: Any, txt: str) -> None:
        t = r._r.find(qn("w:t"))
        if t is None:
            t = etree.SubElement(r._r, qn("w:t"))
        t.text = txt
        t.set(XML_SPACE, "preserve")

    def _fmt_key(r: Any) -> str:
        rpr = r._r.find(qn("w:rPr"))
        return etree.tostring(rpr, encoding="unicode") if rpr is not None else ""

    runs = para.runs
    if not runs:
        para.add_run(new_text)
        return
    if len(runs) == 1:
        _set_run_text(runs[0], new_text)
        return

    fmt_keys = [_fmt_key(r) for r in runs]
    if len(set(fmt_keys)) == 1:
        _set_run_text(runs[0], new_text)
        for run in runs[1:]:
            _set_run_text(run, "")
        return

    orig_texts = [r.text or "" for r in runs]
    distributed: list = [""] * len(runs)
    remaining = new_text
    matched_any = False
    for i, orig in enumerate(orig_texts):
        stripped = orig.strip()
        if not stripped:
            continue
        idx = remaining.lower().find(stripped.lower())
        if idx != -1:
            if i > 0 and not distributed[i - 1] and idx > 0:
                distributed[i - 1] = remaining[:idx].rstrip()
            distributed[i] = remaining[idx: idx + len(stripped)]
            remaining = remaining[idx + len(stripped):].lstrip()
            matched_any = True

    if matched_any:
        if remaining:
            for i in range(len(runs) - 1, -1, -1):
                if distributed[i]:
                    distributed[i] = distributed[i] + " " + remaining.lstrip()
                    break
            else:
                distributed[-1] = remaining
        for i, run in enumerate(runs):
            _set_run_text(run, distributed[i])
        return

    orig_total = sum(len(t) for t in orig_texts) or 1
    new_total = len(new_text)
    words = new_text.split()
    word_idx = 0
    for i, run in enumerate(runs):
        if i == len(runs) - 1:
            chunk = " ".join(words[word_idx:])
        else:
            ratio = len(orig_texts[i]) / orig_total
            target_len = max(1, round(new_total * ratio))
            chunk_words: list = []
            chunk_len = 0
            while word_idx < len(words):
                w = words[word_idx]
                candidate = chunk_len + len(w) + (1 if chunk_words else 0)
                if candidate > target_len and chunk_words:
                    break
                chunk_words.append(w)
                chunk_len = candidate
                word_idx += 1
            chunk = " ".join(chunk_words)
        _set_run_text(run, chunk)


def _replace_cell_text_keep_fmt(cell: Any, new_text: str) -> None:
    """Thay text của table cell, giữ formatting."""
    try:
        from docx.oxml.ns import qn
    except ImportError:
        cell.text = new_text
        return
    paras = cell.paragraphs
    if not paras:
        cell.text = new_text
        return
    _replace_para_text_keep_fmt(paras[0], new_text)
    tc = cell._tc
    all_p = tc.findall(qn("w:p"))
    for extra_p in all_p[1:]:
        tc.remove(extra_p)


def build_translated_docx(
    content: bytes,
    columns: List[str],
    rows: List[dict],
    to_translate: List[str],
) -> bytes:
    """Tạo lại file DOCX giữ cấu trúc và formatting, áp dụng bản dịch từ rows."""
    try:
        from docx import Document
    except ImportError:
        return content
    if not content or len(content) < 4:
        return content
    try:
        doc = Document(io.BytesIO(content))
    except Exception:
        return content
    blocks = _collect_docx_blocks(doc, max_items=len(rows))
    for i, row_data in enumerate(rows):
        if i >= len(blocks):
            break
        val = row_data.get("Nội dung_translated")
        if val is None:
            continue
        val = val.strip()
        _text, kind, ref = blocks[i]
        if kind == "para":
            _replace_para_text_keep_fmt(ref, val)
        else:
            _replace_cell_text_keep_fmt(ref, val)
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def export_docx(columns: List[str], rows: List[dict]) -> bytes:
    """Xuất columns + rows ra file .docx (bảng)."""
    try:
        from docx import Document
    except ImportError:
        raise HTTPException(status_code=503, detail="python-docx chưa được cài đặt.")
    doc = Document()
    table = doc.add_table(rows=1 + len(rows), cols=len(columns))
    for j, col in enumerate(columns):
        table.rows[0].cells[j].text = str(col)
    for i, row in enumerate(rows):
        for j, col in enumerate(columns):
            table.rows[i + 1].cells[j].text = str(row.get(col, "") or "")
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
